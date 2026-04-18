"""
Convert a DoxDocument to a Microsoft Word (.docx) file.

Handles all .dox element types including:
  - Headings (H1-H6 mapped to Word heading styles)
  - Paragraphs with inline formatting (bold, italic, code, links)
  - Tables with colspan/rowspan via cell merging
  - Code blocks with monospace font
  - Math blocks (rendered as styled text — Word's OMML not used)
  - Figures (embedded if file exists, placeholder otherwise)
  - Lists (ordered and unordered)
  - Page breaks (inserted between pages)
  - Footnotes
  - Form fields (rendered as text)

Usage:
    from dox.converters.to_docx import to_docx, to_docx_bytes
    to_docx(doc, "output.docx")
    raw_bytes = to_docx_bytes(doc)
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Any

from dox.converters._figure_utils import figure_binary_source
from dox.models.document import DoxDocument

logger = logging.getLogger(__name__)

# Compile regex pattern once at module level for reuse
_INLINE_FORMAT_PATTERN = re.compile(
    r'(\*\*(.+?)\*\*)'   # bold
    r'|(\*(.+?)\*)'      # italic
    r'|(`(.+?)`)'        # inline code
    r'|(\[([^\]]+)\]\(([^)]+)\))'  # link
)
from dox.models.elements import (
    Annotation,
    Blockquote,
    Chart,
    CodeBlock,
    CrossRef,
    Element,
    Figure,
    Footnote,
    FormField,
    Heading,
    HorizontalRule,
    KeyValuePair,
    ListBlock,
    MathBlock,
    PageBreak,
    Paragraph,
    Table,
)


def to_docx(
    doc: DoxDocument,
    output_path: str | Path,
    *,
    page_width_inches: float = 8.5,
    page_height_inches: float = 11.0,
    margin_inches: float = 1.0,
) -> Path:
    """
    Convert a DoxDocument to a .docx file.

    Args:
        doc: The DoxDocument to convert.
        output_path: File path for the output .docx.
        page_width_inches: Page width.
        page_height_inches: Page height.
        margin_inches: Page margins (all sides).

    Returns:
        Path to the created .docx file.
    """
    try:
        from docx import Document as WordDocument
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
    except ImportError:
        raise ImportError(
            "python-docx is required: pip install python-docx"
        )

    output_path = Path(output_path)
    word_doc = WordDocument()

    # Page setup
    section = word_doc.sections[0]
    section.page_width = Inches(page_width_inches)
    section.page_height = Inches(page_height_inches)
    section.top_margin = Inches(margin_inches)
    section.bottom_margin = Inches(margin_inches)
    section.left_margin = Inches(margin_inches)
    section.right_margin = Inches(margin_inches)

    # Process elements
    for element in doc.elements:
        _add_element(word_doc, element)

    word_doc.save(str(output_path))
    return output_path


def to_docx_bytes(doc: DoxDocument) -> bytes:
    """
    Convert a DoxDocument to .docx bytes in memory.

    Returns:
        Raw .docx file bytes.
    """
    try:
        from docx import Document as WordDocument
        from docx.shared import Inches
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    word_doc = WordDocument()

    section = word_doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    for element in doc.elements:
        _add_element(word_doc, element)

    buf = io.BytesIO()
    try:
        word_doc.save(buf)
        return buf.getvalue()
    finally:
        buf.close()


def _add_element(word_doc: Any, element: Element) -> None:
    """Dispatch element to the appropriate handler."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if isinstance(element, PageBreak):
        _add_page_break(word_doc)

    elif isinstance(element, HorizontalRule):
        _add_horizontal_rule(word_doc)

    elif isinstance(element, Heading):
        _add_heading(word_doc, element)

    elif isinstance(element, Paragraph):
        _add_paragraph(word_doc, element)

    elif isinstance(element, Blockquote):
        _add_blockquote(word_doc, element)

    elif isinstance(element, Table):
        _add_table(word_doc, element)

    elif isinstance(element, CodeBlock):
        _add_code_block(word_doc, element)

    elif isinstance(element, MathBlock):
        _add_math_block(word_doc, element)

    elif isinstance(element, Figure):
        _add_figure(word_doc, element)

    elif isinstance(element, ListBlock):
        _add_list(word_doc, element)

    elif isinstance(element, Footnote):
        _add_footnote(word_doc, element)

    elif isinstance(element, FormField):
        _add_form_field(word_doc, element)

    elif isinstance(element, Chart):
        _add_chart_placeholder(word_doc, element)

    elif isinstance(element, Annotation):
        _add_annotation(word_doc, element)

    elif isinstance(element, KeyValuePair):
        _add_kv_pair(word_doc, element)

    elif isinstance(element, CrossRef):
        _add_crossref(word_doc, element)


# ------------------------------------------------------------------
# Element handlers
# ------------------------------------------------------------------

def _add_page_break(word_doc: Any) -> None:
    from docx.enum.text import WD_BREAK
    p = word_doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_horizontal_rule(word_doc: Any) -> None:
    """Add a horizontal line using a bottom border."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = word_doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')  # 3pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')

    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_heading(word_doc: Any, element: Heading) -> None:
    # Word supports heading levels 0-9, .dox uses 1-6
    level = min(max(element.level, 1), 9)
    word_doc.add_heading(element.text or "", level=level)


def _add_paragraph(word_doc: Any, element: Paragraph) -> None:
    p = word_doc.add_paragraph()
    _add_inline_text(p, element.text or "")


def _add_blockquote(word_doc: Any, element: Blockquote) -> None:
    """Add blockquote using Word's built-in Quote style."""
    from docx.shared import Inches
    p = word_doc.add_paragraph(element.text or "", style="Intense Quote")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)


def _add_table(word_doc: Any, element: Table) -> None:
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn

    if not element.rows:
        return

    num_rows = element.num_rows
    num_cols = element.num_cols

    if num_rows == 0 or num_cols == 0:
        return

    # Caption above table
    if element.caption:
        cap_p = word_doc.add_paragraph()
        run = cap_p.add_run(element.caption)
        run.italic = True

    # Create table
    table = word_doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    # Track merged cells to avoid overwriting
    merged_cells: set[tuple[int, int]] = set()

    for r_idx, row in enumerate(element.rows):
        semantic_c_idx = 0
        for cell in row.cells:
            c_idx = semantic_c_idx
            if (r_idx, c_idx) in merged_cells:
                semantic_c_idx += max(1, cell.colspan)
                continue

            word_cell = table.cell(r_idx, c_idx)

            # Handle colspan and rowspan via cell merging
            end_r = r_idx + cell.rowspan - 1
            end_c = c_idx + cell.colspan - 1

            if cell.colspan > 1 or cell.rowspan > 1:
                # Clamp to table bounds
                end_r = min(end_r, num_rows - 1)
                end_c = min(end_c, num_cols - 1)

                try:
                    merge_cell = table.cell(end_r, end_c)
                    word_cell.merge(merge_cell)
                except (ValueError, IndexError):
                    # Merge can fail if cells already merged or out of bounds
                    logger.debug(f"Cell merge failed for ({r_idx}, {c_idx}) to ({end_r}, {end_c})")

                # Mark spanned cells
                for mr in range(r_idx, end_r + 1):
                    for mc in range(c_idx, end_c + 1):
                        if (mr, mc) != (r_idx, c_idx):
                            merged_cells.add((mr, mc))

            # Set cell text
            word_cell.text = cell.text

            # Header styling
            if cell.is_header or row.is_header:
                for paragraph in word_cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            semantic_c_idx += max(1, cell.colspan)

    # Auto-fit table width
    try:
        table.autofit = True
    except Exception:
        pass


def _add_code_block(word_doc: Any, element: CodeBlock) -> None:
    from docx.shared import Pt, RGBColor

    p = word_doc.add_paragraph()
    run = p.add_run(element.code or "")
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Light gray background via paragraph shading
    _set_paragraph_shading(p, "F5F5F5")


def _add_math_block(word_doc: Any, element: MathBlock) -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = word_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(element.expression or "")
    run.italic = True
    run.font.size = Pt(11)


def _add_figure(word_doc: Any, element: Figure) -> None:
    from docx.shared import Inches

    image_source, label = figure_binary_source(element)
    if image_source is not None:
        try:
            word_doc.add_picture(image_source, width=Inches(5))
        except Exception:
            # Fallback to placeholder
            p = word_doc.add_paragraph()
            run = p.add_run(f"[Image: {label}]")
            run.italic = True
    else:
        p = word_doc.add_paragraph()
        run = p.add_run(f"[Image: {label}]")
        run.italic = True

    # Caption
    if element.caption:
        cap_p = word_doc.add_paragraph()
        run = cap_p.add_run(element.caption)
        run.italic = True
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_list(word_doc: Any, element: ListBlock) -> None:
    for idx, item in enumerate(element.items):
        if element.ordered:
            style = 'List Number'
        else:
            style = 'List Bullet'
        p = word_doc.add_paragraph(style=style)
        _add_inline_text(p, item.text)
        # Handle nested items
        for child in item.children:
            child_style = 'List Number 2' if element.ordered else 'List Bullet 2'
            child_p = word_doc.add_paragraph(style=child_style)
            _add_inline_text(child_p, child.text)


def _add_footnote(word_doc: Any, element: Footnote) -> None:
    from docx.shared import Pt

    p = word_doc.add_paragraph()
    # Superscript number
    num_run = p.add_run(f"{element.number} ")
    num_run.font.superscript = True
    num_run.font.size = Pt(8)
    # Footnote text
    text_run = p.add_run(element.text)
    text_run.font.size = Pt(9)


def _add_form_field(word_doc: Any, element: FormField) -> None:
    p = word_doc.add_paragraph()
    label_run = p.add_run(f"{element.field_name}: ")
    label_run.bold = True
    p.add_run(element.value)


def _add_chart_placeholder(word_doc: Any, element: Chart) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = word_doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[Chart: {element.chart_type}]")
    run.italic = True


def _add_annotation(word_doc: Any, element: Annotation) -> None:
    from docx.shared import RGBColor

    p = word_doc.add_paragraph()
    run = p.add_run(f"[{element.annotation_type}] {element.text}")
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x00)


def _add_kv_pair(word_doc: Any, element: KeyValuePair) -> None:
    p = word_doc.add_paragraph()
    key_run = p.add_run(f"{element.key}: ")
    key_run.bold = True
    p.add_run(element.value)


def _add_crossref(word_doc: Any, element: CrossRef) -> None:
    p = word_doc.add_paragraph()
    run = p.add_run(f"[→ {element.ref_type}:{element.ref_id}]")
    run.italic = True


# ------------------------------------------------------------------
# Inline formatting
# ------------------------------------------------------------------

def _add_inline_text(paragraph: Any, text: str) -> None:
    """Parse basic Markdown inline formatting and add as Word runs."""
    from docx.shared import RGBColor

    last_end = 0
    for m in _INLINE_FORMAT_PATTERN.finditer(text):
        # Add text before this match
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])

        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(4):  # italic
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(6):  # code
            run = paragraph.add_run(m.group(6))
            run.font.name = 'Courier New'
        elif m.group(8):  # link
            # Word hyperlinks are complex — render as underlined text
            link_text = m.group(8)
            run = paragraph.add_run(link_text)
            run.underline = True
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

        last_end = m.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _set_paragraph_shading(paragraph: Any, color_hex: str) -> None:
    """Set paragraph background color."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(
        paragraph.paragraph_format.element.get_or_add_pPr(),
        qn('w:shd')
    )
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
