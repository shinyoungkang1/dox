"""
Round-trip conversion tests: .dox → DOCX/PDF → re-extract → compare.

Verifies that the full pipeline preserves content:
  1. Parse .dox → DoxDocument
  2. Convert DoxDocument → DOCX or PDF
  3. Re-extract text from the output file
  4. Compare extracted text against original elements

Also tests:
  - Serialize → Parse round-trip (lossless within format)
  - All element types survive DOCX conversion
  - All element types survive PDF conversion
  - Table structure (including colspan/rowspan) preserved in DOCX
  - Inline formatting (bold, italic, code, links) preserved
"""

from __future__ import annotations

import re
import tempfile
from difflib import SequenceMatcher
from pathlib import Path

import pytest

from dox.converters.to_docx import to_docx, to_docx_bytes
from dox.converters.to_pdf import to_pdf, to_pdf_bytes
from dox.models.document import DoxDocument, Frontmatter
from dox.models.elements import (
    Annotation,
    Chart,
    CodeBlock,
    CrossRef,
    Figure,
    Footnote,
    FormField,
    Heading,
    ListBlock,
    ListItem,
    MathBlock,
    PageBreak,
    Paragraph,
    Table,
    TableCell,
    TableRow,
)
from dox.parsers.parser import DoxParser
from dox.serializer import DoxSerializer


# ------------------------------------------------------------------
# Helpers
# ------------------------------------------------------------------

def _build_full_doc() -> DoxDocument:
    """Build a DoxDocument with every element type for testing."""
    doc = DoxDocument()
    doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")

    doc.add_element(Heading(level=1, text="Main Title", page=1))
    doc.add_element(Paragraph(
        text="This is a **bold** and *italic* paragraph with `code` inline.",
        page=1,
    ))
    doc.add_element(Heading(level=2, text="Tables Section", page=1))

    # Simple table
    t = Table(caption="Simple Table", table_id="t1")
    t.rows.append(TableRow(
        cells=[TableCell(text="Name", is_header=True),
               TableCell(text="Value", is_header=True)],
        is_header=True,
    ))
    t.rows.append(TableRow(
        cells=[TableCell(text="Alpha"), TableCell(text="100")],
    ))
    t.rows.append(TableRow(
        cells=[TableCell(text="Beta"), TableCell(text="200")],
    ))
    doc.add_element(t)

    # Table with colspan
    t2 = Table(caption="Span Table", table_id="t2")
    t2.rows.append(TableRow(
        cells=[TableCell(text="Header", is_header=True, colspan=2),
               TableCell(text="H2", is_header=True)],
        is_header=True,
    ))
    t2.rows.append(TableRow(
        cells=[TableCell(text="A"), TableCell(text="B"), TableCell(text="C")],
    ))
    t2.rows.append(TableRow(
        cells=[TableCell(text="D", rowspan=2), TableCell(text="E"),
               TableCell(text="F")],
    ))
    t2.rows.append(TableRow(
        cells=[TableCell(text=""), TableCell(text="G"), TableCell(text="H")],
    ))
    doc.add_element(t2)

    doc.add_element(Heading(level=2, text="Code Section", page=2))
    doc.add_element(CodeBlock(
        code="def hello():\n    print('world')",
        language="python",
        page=2,
    ))

    doc.add_element(MathBlock(
        expression="E = mc^2",
        display_mode=True,
        page=2,
    ))

    doc.add_element(Figure(
        source="diagram.png",
        caption="Architecture Diagram",
        figure_id="fig1",
        page=2,
    ))

    doc.add_element(ListBlock(
        items=[ListItem(text="First item"),
               ListItem(text="Second item"),
               ListItem(text="Third item")],
        ordered=False,
    ))

    doc.add_element(ListBlock(
        items=[ListItem(text="Step one"),
               ListItem(text="Step two")],
        ordered=True,
    ))

    doc.add_element(Footnote(number=1, text="This is a footnote reference.", page=3))
    doc.add_element(FormField(field_name="Author", value="John Doe"))
    doc.add_element(PageBreak(from_page=3, to_page=4))
    doc.add_element(Heading(level=1, text="Page Four Content", page=4))
    doc.add_element(Paragraph(text="Final paragraph on page four.", page=4))

    return doc


def _extract_docx_text(path: Path) -> list[str]:
    """Extract all text content from a DOCX file."""
    from docx import Document as WordDocument
    wd = WordDocument(str(path))
    texts = []
    for p in wd.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())
    for table in wd.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())
    return texts


def _extract_pdf_text(path: Path) -> list[str]:
    """Extract all text content from a PDF file."""
    import fitz
    texts = []
    pdf = fitz.open(str(path))
    for page in pdf:
        text = page.get_text()
        for line in text.split('\n'):
            if line.strip():
                texts.append(line.strip())
    pdf.close()
    return texts


def _text_similarity(a: str, b: str) -> float:
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()


def _find_in_texts(needle: str, texts: list[str], threshold: float = 0.7) -> bool:
    """Check if a string appears (approximately) in a list of text lines."""
    needle_lower = needle.lower()
    for t in texts:
        if needle_lower in t.lower():
            return True
        if _text_similarity(needle, t) >= threshold:
            return True
    return False


# ------------------------------------------------------------------
# .dox serialize → parse round-trip
# ------------------------------------------------------------------

class TestDoxSerializeRoundtrip:
    """Verify .dox → text → .dox is lossless."""

    def test_heading_roundtrip(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")
        doc.add_element(Heading(level=2, text="Hello World", page=5, element_id="h2"))

        serializer = DoxSerializer()
        text = serializer.serialize(doc, include_spatial=False, include_metadata=False)

        parser = DoxParser()
        doc2 = parser.parse(text)

        assert len(doc2.elements) == 1
        h = doc2.elements[0]
        assert isinstance(h, Heading)
        assert h.level == 2
        assert h.text == "Hello World"
        assert h.page == 5
        assert h.element_id == "h2"

    def test_paragraph_roundtrip(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")
        doc.add_element(Paragraph(
            text="A **bold** and *italic* paragraph.",
            page=1,
            confidence=0.95,
        ))

        serializer = DoxSerializer()
        text = serializer.serialize(doc, include_spatial=False, include_metadata=False)

        parser = DoxParser()
        doc2 = parser.parse(text)

        assert len(doc2.elements) == 1
        p = doc2.elements[0]
        assert isinstance(p, Paragraph)
        assert "bold" in p.text
        assert "italic" in p.text
        assert p.page == 1
        assert p.confidence == pytest.approx(0.95, abs=0.01)

    def test_table_roundtrip_with_spans(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")

        t = Table(table_id="t1", caption="Test")
        t.rows.append(TableRow(
            cells=[TableCell(text="Merged", is_header=True, colspan=2)],
            is_header=True,
        ))
        t.rows.append(TableRow(
            cells=[TableCell(text="A"), TableCell(text="B")],
        ))
        doc.add_element(t)

        serializer = DoxSerializer()
        text = serializer.serialize(doc, include_spatial=False, include_metadata=False)

        parser = DoxParser()
        doc2 = parser.parse(text)

        tables = [e for e in doc2.elements if isinstance(e, Table)]
        assert len(tables) == 1
        assert tables[0].rows[0].cells[0].colspan == 2
        assert tables[0].rows[0].cells[0].text == "Merged"

    def test_math_roundtrip(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")
        doc.add_element(MathBlock(
            expression="\\int_0^\\infty e^{-x} dx = 1",
            display_mode=True,
            page=3,
        ))

        serializer = DoxSerializer()
        text = serializer.serialize(doc, include_spatial=False, include_metadata=False)

        parser = DoxParser()
        doc2 = parser.parse(text)

        assert len(doc2.elements) == 1
        m = doc2.elements[0]
        assert isinstance(m, MathBlock)
        assert "\\int" in m.expression
        assert m.page == 3

    def test_code_block_roundtrip(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")
        doc.add_element(CodeBlock(
            code="for i in range(10):\n    print(i)",
            language="python",
        ))

        serializer = DoxSerializer()
        text = serializer.serialize(doc, include_spatial=False, include_metadata=False)

        parser = DoxParser()
        doc2 = parser.parse(text)

        codes = [e for e in doc2.elements if isinstance(e, CodeBlock)]
        assert len(codes) == 1
        assert "range(10)" in codes[0].code
        assert codes[0].language == "python"

    def test_list_roundtrip(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")
        doc.add_element(ListBlock(
            items=[ListItem(text="Alpha"), ListItem(text="Beta"), ListItem(text="Gamma")],
            ordered=True,
            start=1,
        ))

        serializer = DoxSerializer()
        text = serializer.serialize(doc, include_spatial=False, include_metadata=False)

        parser = DoxParser()
        doc2 = parser.parse(text)

        lists = [e for e in doc2.elements if isinstance(e, ListBlock)]
        assert len(lists) == 1
        assert lists[0].ordered is True
        assert len(lists[0].items) == 3
        assert lists[0].items[1].text == "Beta"

    def test_full_doc_roundtrip(self):
        """Full document text content survives serialize→parse."""
        doc = _build_full_doc()

        serializer = DoxSerializer()
        text = serializer.serialize(doc, include_spatial=False, include_metadata=False)

        parser = DoxParser()
        doc2 = parser.parse(text)

        # Collect all text from both versions
        def _all_text(d):
            texts = []
            for el in d.elements:
                for attr in ("text", "code", "expression", "field_name"):
                    val = getattr(el, attr, None)
                    if val and isinstance(val, str):
                        texts.append(val)
                if hasattr(el, "items"):
                    for item in el.items:
                        texts.append(item.text)
                if hasattr(el, "rows"):
                    for row in el.rows:
                        for cell in row.cells:
                            texts.append(cell.text)
            return " ".join(texts).lower()

        rt_text = _all_text(doc2)

        # Key content should survive
        for keyword in ["main title", "bold", "italic", "alpha", "100",
                         "hello", "mc^2", "first item", "step one", "footnote"]:
            assert keyword in rt_text, f"Lost content: {keyword}"

        # Should have a reasonable number of elements (parser may merge/reclassify some)
        assert len(doc2.elements) >= 5


# ------------------------------------------------------------------
# .dox → DOCX round-trip
# ------------------------------------------------------------------

class TestDocxRoundtrip:
    """Verify content survives .dox → DOCX conversion."""

    def test_basic_docx_generation(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.exists()
        assert path.stat().st_size > 0

    def test_docx_bytes(self):
        doc = _build_full_doc()
        raw = to_docx_bytes(doc)
        assert len(raw) > 1000
        # DOCX is a ZIP file — check magic bytes
        assert raw[:2] == b"PK"

    def test_headings_in_docx(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        texts = _extract_docx_text(path)
        assert _find_in_texts("Main Title", texts)
        assert _find_in_texts("Tables Section", texts)
        assert _find_in_texts("Code Section", texts)
        assert _find_in_texts("Page Four Content", texts)

    def test_paragraph_text_in_docx(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        texts = _extract_docx_text(path)
        assert _find_in_texts("bold", texts)
        assert _find_in_texts("italic", texts)
        assert _find_in_texts("Final paragraph", texts)

    def test_table_data_in_docx(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        from docx import Document as WordDocument
        wd = WordDocument(str(path))

        assert len(wd.tables) >= 2

        # Check first table content
        t1 = wd.tables[0]
        all_text = [c.text for r in t1.rows for c in r.cells]
        assert "Name" in all_text
        assert "Alpha" in all_text
        assert "100" in all_text

    def test_table_colspan_in_docx(self):
        """Verify that colspan merging works in DOCX output."""
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        from docx import Document as WordDocument
        wd = WordDocument(str(path))

        # Second table has colspan=2 header
        t2 = wd.tables[1]
        # After merge, the first row should have merged cells
        # In python-docx, merged cells report the same text
        first_row_texts = [c.text for c in t2.rows[0].cells]
        # "Header" should appear (colspan=2 merged)
        assert "Header" in first_row_texts

    def test_code_block_in_docx(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        texts = _extract_docx_text(path)
        assert _find_in_texts("def hello", texts)
        assert _find_in_texts("print", texts)

    def test_math_in_docx(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        texts = _extract_docx_text(path)
        assert _find_in_texts("E = mc^2", texts)

    def test_list_items_in_docx(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        texts = _extract_docx_text(path)
        assert _find_in_texts("First item", texts)
        assert _find_in_texts("Second item", texts)
        assert _find_in_texts("Step one", texts)

    def test_footnote_in_docx(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        texts = _extract_docx_text(path)
        assert _find_in_texts("footnote reference", texts)

    def test_form_field_in_docx(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        texts = _extract_docx_text(path)
        assert _find_in_texts("Author", texts)
        assert _find_in_texts("John Doe", texts)

    def test_inline_formatting_preserved(self):
        """Verify bold/italic/code runs exist in DOCX."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")
        doc.add_element(Paragraph(text="Text with **bold** and *italic* and `code` words."))

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        from docx import Document as WordDocument
        wd = WordDocument(str(path))

        # Check that runs with bold/italic/monospace exist
        has_bold = False
        has_italic = False
        has_mono = False
        for p in wd.paragraphs:
            for run in p.runs:
                if run.bold:
                    has_bold = True
                if run.italic:
                    has_italic = True
                if run.font.name == 'Courier New':
                    has_mono = True

        assert has_bold, "Bold formatting lost"
        assert has_italic, "Italic formatting lost"
        assert has_mono, "Code/monospace formatting lost"


# ------------------------------------------------------------------
# .dox → PDF round-trip
# ------------------------------------------------------------------

class TestPdfRoundtrip:
    """Verify content survives .dox → PDF conversion."""

    def test_basic_pdf_generation(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)
        assert path.exists()
        assert path.stat().st_size > 0

    def test_pdf_bytes(self):
        doc = _build_full_doc()
        raw = to_pdf_bytes(doc)
        assert len(raw) > 500
        assert raw[:5] == b"%PDF-"

    def test_headings_in_pdf(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        assert _find_in_texts("Main Title", texts)
        assert _find_in_texts("Tables Section", texts)
        assert _find_in_texts("Code Section", texts)

    def test_paragraph_text_in_pdf(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        assert _find_in_texts("bold", texts)
        assert _find_in_texts("italic", texts)

    def test_table_data_in_pdf(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        assert _find_in_texts("Alpha", texts)
        assert _find_in_texts("100", texts)
        assert _find_in_texts("Beta", texts)

    def test_code_in_pdf(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        assert _find_in_texts("hello", texts)

    def test_math_in_pdf(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        assert _find_in_texts("E = mc^2", texts)

    def test_list_items_in_pdf(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        assert _find_in_texts("First item", texts)
        assert _find_in_texts("Step one", texts)

    def test_footnote_in_pdf(self):
        doc = _build_full_doc()
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        assert _find_in_texts("footnote", texts)


# ------------------------------------------------------------------
# Full pipeline: .dox text → parse → DOCX → re-extract → compare
# ------------------------------------------------------------------

class TestFullPipelineDocx:
    """End-to-end: .dox source → DoxDocument → DOCX → text extraction → comparison."""

    DOX_SOURCE = """---dox
version: '1.0'
source: pipeline-test
lang: en
---

# Quarterly Report

## Executive Summary

The company achieved **record revenue** of $2.5M in Q3, representing a *15% increase*
over the previous quarter. Key drivers include `product-x` launch and expanded
market presence.

||| table caption="Q3 Financial Summary"
| Metric         | Q2          | Q3          | Change  |
|----------------|-------------|-------------|---------|
| Revenue        | $2.17M      | $2.50M      | +15.2%  |
| Expenses       | $1.80M      | $1.95M      | +8.3%   |
| Net Income     | $0.37M      | $0.55M      | +48.6%  |
|||

## Technical Highlights

```python
# Performance improvement
def optimize_query(query):
    return query.with_index('revenue_date_idx')
```

$$\\Delta R = R_{Q3} - R_{Q2} = 0.33M$$

- Reduced latency by 40%
- Improved cache hit rate to 95%
- Zero downtime deployments

1. Phase 1: Infrastructure upgrade
2. Phase 2: Application optimization
3. Phase 3: Monitoring rollout

[^1]: Revenue figures are unaudited.
"""

    def test_pipeline_text_preservation(self):
        """All significant text survives the .dox → DOCX pipeline."""
        parser = DoxParser()
        doc = parser.parse(self.DOX_SOURCE)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        texts = _extract_docx_text(path)
        full_text = " ".join(texts).lower()

        # Key content checks
        assert "quarterly report" in full_text
        assert "executive summary" in full_text
        assert "record revenue" in full_text
        assert "2.5m" in full_text or "2.50m" in full_text
        assert "revenue" in full_text
        assert "expenses" in full_text
        assert "net income" in full_text
        assert "optimize_query" in full_text
        assert "reduced latency" in full_text
        assert "unaudited" in full_text

    def test_pipeline_table_integrity(self):
        """Table data is complete in DOCX output."""
        parser = DoxParser()
        doc = parser.parse(self.DOX_SOURCE)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        from docx import Document as WordDocument
        wd = WordDocument(str(path))

        assert len(wd.tables) >= 1
        t = wd.tables[0]
        all_cells = [c.text for r in t.rows for c in r.cells]
        all_text = " ".join(all_cells).lower()

        assert "revenue" in all_text
        assert "expenses" in all_text
        assert "+15.2%" in " ".join(all_cells) or "15.2" in all_text

    def test_pipeline_element_count(self):
        """Reasonable number of elements survive the pipeline."""
        parser = DoxParser()
        doc = parser.parse(self.DOX_SOURCE)

        # Should have: 2 headings + 2 paragraphs + 1 table + 1 code +
        # 1 math + 2 lists + 1 footnote = ~10+ elements
        assert len(doc.elements) >= 8

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)

        from docx import Document as WordDocument
        wd = WordDocument(str(path))

        # DOCX should have paragraphs + tables
        assert len(wd.paragraphs) >= 8
        assert len(wd.tables) >= 1


# ------------------------------------------------------------------
# Full pipeline: .dox text → parse → PDF → re-extract → compare
# ------------------------------------------------------------------

class TestFullPipelinePdf:
    """End-to-end: .dox source → DoxDocument → PDF → text extraction → comparison."""

    DOX_SOURCE = TestFullPipelineDocx.DOX_SOURCE

    def test_pipeline_text_preservation(self):
        parser = DoxParser()
        doc = parser.parse(self.DOX_SOURCE)

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        full_text = " ".join(texts).lower()

        assert "quarterly report" in full_text
        assert "executive summary" in full_text
        assert "revenue" in full_text
        assert "optimize_query" in full_text

    def test_pipeline_table_data(self):
        parser = DoxParser()
        doc = parser.parse(self.DOX_SOURCE)

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)

        texts = _extract_pdf_text(path)
        full_text = " ".join(texts).lower()

        assert "revenue" in full_text
        assert "expenses" in full_text


# ------------------------------------------------------------------
# Cross-format consistency
# ------------------------------------------------------------------

class TestCrossFormatConsistency:
    """Verify that DOCX and PDF produce consistent content."""

    def test_same_content_both_formats(self):
        """DOCX and PDF should contain the same core text."""
        doc = _build_full_doc()

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = to_docx(doc, f.name)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            pdf_path = to_pdf(doc, f.name)

        docx_texts = _extract_docx_text(docx_path)
        pdf_texts = _extract_pdf_text(pdf_path)

        docx_full = " ".join(docx_texts).lower()
        pdf_full = " ".join(pdf_texts).lower()

        # Both should contain the key content
        for keyword in ["main title", "alpha", "100", "hello", "first item"]:
            assert keyword in docx_full, f"DOCX missing: {keyword}"
            assert keyword in pdf_full, f"PDF missing: {keyword}"

    def test_text_coverage_score(self):
        """Measure what % of original element text appears in outputs."""
        doc = _build_full_doc()

        # Collect all text from original elements
        original_texts = []
        for el in doc.elements:
            if isinstance(el, Heading):
                original_texts.append(el.text)
            elif isinstance(el, Paragraph):
                # Strip markdown formatting for comparison
                clean = re.sub(r'[*`\[\]()]', '', el.text)
                original_texts.append(clean)
            elif isinstance(el, CodeBlock):
                original_texts.append(el.code.split('\n')[0])  # First line
            elif isinstance(el, MathBlock):
                original_texts.append(el.expression)
            elif isinstance(el, ListBlock):
                for item in el.items:
                    original_texts.append(item.text)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = to_docx(doc, f.name)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            pdf_path = to_pdf(doc, f.name)

        docx_texts = _extract_docx_text(docx_path)
        pdf_texts = _extract_pdf_text(pdf_path)

        docx_found = sum(1 for t in original_texts if _find_in_texts(t, docx_texts))
        pdf_found = sum(1 for t in original_texts if _find_in_texts(t, pdf_texts))

        docx_coverage = docx_found / len(original_texts) if original_texts else 0
        pdf_coverage = pdf_found / len(original_texts) if original_texts else 0

        print(f"\nText coverage — DOCX: {docx_coverage:.1%} ({docx_found}/{len(original_texts)})")
        print(f"Text coverage — PDF:  {pdf_coverage:.1%} ({pdf_found}/{len(original_texts)})")

        assert docx_coverage >= 0.85, f"DOCX coverage too low: {docx_coverage:.1%}"
        assert pdf_coverage >= 0.80, f"PDF coverage too low: {pdf_coverage:.1%}"


# ------------------------------------------------------------------
# Edge cases
# ------------------------------------------------------------------

class TestEdgeCases:
    """Edge case handling for converters."""

    def test_empty_document_docx(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="empty", lang="en")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.exists()

    def test_empty_document_pdf(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="empty", lang="en")
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)
        assert path.exists()

    def test_unicode_content(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")
        doc.add_element(Heading(level=1, text="日本語タイトル"))
        doc.add_element(Paragraph(text="Ünïcödë tëxt with émojis and spëcial chars: <>&"))

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = to_docx(doc, f.name)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            pdf_path = to_pdf(doc, f.name)

        assert docx_path.exists()
        assert pdf_path.exists()

    def test_very_long_table(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")

        t = Table()
        t.rows.append(TableRow(
            cells=[TableCell(text="ID", is_header=True),
                   TableCell(text="Value", is_header=True)],
            is_header=True,
        ))
        for i in range(50):
            t.rows.append(TableRow(
                cells=[TableCell(text=str(i)), TableCell(text=f"val_{i}")],
            ))
        doc.add_element(t)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = to_docx(doc, f.name)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            pdf_path = to_pdf(doc, f.name)

        # Both should have all 50 data rows
        docx_texts = _extract_docx_text(docx_path)
        assert _find_in_texts("val_49", docx_texts)

        pdf_texts = _extract_pdf_text(pdf_path)
        assert _find_in_texts("val_49", pdf_texts)

    def test_special_chars_in_table(self):
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="test", lang="en")

        t = Table()
        t.rows.append(TableRow(
            cells=[TableCell(text="<html>", is_header=True),
                   TableCell(text="A & B", is_header=True)],
            is_header=True,
        ))
        t.rows.append(TableRow(
            cells=[TableCell(text="100%"), TableCell(text="$5.00")],
        ))
        doc.add_element(t)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = to_docx(doc, f.name)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            pdf_path = to_pdf(doc, f.name)

        assert docx_path.exists()
        assert pdf_path.exists()
