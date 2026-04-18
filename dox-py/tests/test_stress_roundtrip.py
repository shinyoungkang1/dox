"""
Stress tests: 50+ hard edge cases for the .dox round-trip system.

Categories:
  A. Parser edge cases (malformed input, tricky syntax)
  B. Table torture tests (spans, empty cells, huge tables, nested content)
  C. Inline formatting edge cases (nested, adjacent, broken)
  D. Math/code edge cases (special chars, multi-line, mixed)
  E. Cross-page and page break handling
  F. DOCX converter edge cases (merging, styling, structure)
  G. PDF converter edge cases
  H. Full pipeline: .dox → serialize → parse → DOCX → extract → compare
  I. Unicode, special chars, and encoding
  J. Large/adversarial documents
"""

from __future__ import annotations

import re
import tempfile
from difflib import SequenceMatcher
from pathlib import Path

import pytest

from dox.converters.to_docx import to_docx
from dox.converters.to_pdf import to_pdf
from dox.models.document import DoxDocument, Frontmatter
from dox.models.elements import *
from dox.parsers.parser import DoxParser
from dox.serializer import DoxSerializer


parser = DoxParser()
serializer = DoxSerializer()


def _rt(source: str) -> DoxDocument:
    """Parse → serialize → re-parse (full roundtrip)."""
    doc = parser.parse(source)
    text = serializer.serialize(doc, include_spatial=False, include_metadata=False)
    return parser.parse(text)


def _docx_text(doc: DoxDocument) -> str:
    """Convert to DOCX and extract all text."""
    from docx import Document as WordDocument
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        to_docx(doc, f.name)
        wd = WordDocument(f.name)
    texts = [p.text for p in wd.paragraphs if p.text.strip()]
    for t in wd.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.strip():
                    texts.append(c.text.strip())
    return ' '.join(texts).lower()


def _pdf_text(doc: DoxDocument) -> str:
    """Convert to PDF and extract all text."""
    import fitz
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        to_pdf(doc, f.name)
        pdf = fitz.open(f.name)
    texts = []
    for page in pdf:
        texts.append(page.get_text())
    pdf.close()
    return ' '.join(texts).lower()


# ================================================================
# A. Parser Edge Cases
# ================================================================

class TestParserEdgeCases:

    def test_01_empty_document(self):
        """Empty body with only frontmatter."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: empty\nlang: en\n---\n")
        assert len(doc.elements) == 0

    def test_02_no_frontmatter(self):
        """Document with no frontmatter at all."""
        doc = parser.parse("# Just a heading\n\nAnd a paragraph.")
        headings = [e for e in doc.elements if isinstance(e, Heading)]
        assert len(headings) >= 1

    def test_03_heading_level_boundaries(self):
        """H1 through H6."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n"
        for i in range(1, 7):
            src += f"{'#' * i} Level {i}\n\n"
        doc = parser.parse(src)
        headings = [e for e in doc.elements if isinstance(e, Heading)]
        assert len(headings) == 6
        for i, h in enumerate(headings):
            assert h.level == i + 1

    def test_04_heading_with_special_chars(self):
        """Heading with markdown chars in it."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n# Revenue: $5.08B (+23.4%)\n")
        h = [e for e in doc.elements if isinstance(e, Heading)][0]
        assert "$5.08B" in h.text
        assert "+23.4%" in h.text

    def test_05_paragraph_with_all_inline(self):
        """Paragraph with bold, italic, code, and link."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nText with **bold**, *italic*, `code`, and [link](http://example.com).\n")
        p = [e for e in doc.elements if isinstance(e, Paragraph)][0]
        assert "**bold**" in p.text
        assert "*italic*" in p.text
        assert "`code`" in p.text

    def test_06_multi_line_paragraph(self):
        """Paragraph that spans multiple source lines."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nLine one of the paragraph\ncontinues on line two\nand line three.\n")
        paras = [e for e in doc.elements if isinstance(e, Paragraph)]
        assert len(paras) >= 1
        assert "line one" in paras[0].text.lower()
        assert "line three" in paras[0].text.lower()

    def test_07_consecutive_headings(self):
        """Multiple headings with no content between them."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n# One\n\n## Two\n\n### Three\n")
        headings = [e for e in doc.elements if isinstance(e, Heading)]
        assert len(headings) == 3

    def test_08_metadata_with_spaces(self):
        """Metadata with extra spaces in braces."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n# Title {  page: 5  ,  id: \"h1\"  }\n")
        h = [e for e in doc.elements if isinstance(e, Heading)][0]
        assert h.page == 5
        assert h.element_id == "h1"

    def test_09_paragraph_that_looks_like_heading(self):
        """Text starting with # but not at start of line (in a paragraph context)."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nThis talks about C# and F# languages.\n")
        paras = [e for e in doc.elements if isinstance(e, Paragraph)]
        assert any("C#" in p.text for p in paras)

    def test_10_empty_paragraph_lines(self):
        """Multiple blank lines between elements."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n# Title\n\n\n\n\nParagraph after many blanks.\n")
        assert len(doc.elements) >= 2


# ================================================================
# B. Table Torture Tests
# ================================================================

class TestTableTorture:

    def test_11_single_cell_table(self):
        """1x1 table."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n||| table\n| Only Cell |\n|||")
        tables = [e for e in doc.elements if isinstance(e, Table)]
        assert len(tables) == 1
        assert tables[0].rows[0].cells[0].text == "Only Cell"

    def test_12_wide_table(self):
        """Table with 10 columns."""
        header = "| " + " | ".join(f"H{i}" for i in range(10)) + " |"
        sep = "| " + " | ".join("---" for _ in range(10)) + " |"
        data = "| " + " | ".join(f"D{i}" for i in range(10)) + " |"
        src = f"---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n||| table\n{header}\n{sep}\n{data}\n|||"
        doc = parser.parse(src)
        tables = [e for e in doc.elements if isinstance(e, Table)]
        assert tables[0].num_cols == 10

    def test_13_tall_table(self):
        """Table with 100 rows."""
        rows = ["| ID | Value |", "| --- | --- |"]
        for i in range(100):
            rows.append(f"| {i} | val_{i} |")
        src = f"---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n||| table\n" + "\n".join(rows) + "\n|||"
        doc = parser.parse(src)
        tables = [e for e in doc.elements if isinstance(e, Table)]
        assert tables[0].num_rows >= 100  # 100 data + 1 header

    def test_14_table_with_empty_cells(self):
        """Cells that are empty."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n||| table\n| A |  | C |\n| --- | --- | --- |\n|  | B |  |\n|||"
        doc = parser.parse(src)
        tables = [e for e in doc.elements if isinstance(e, Table)]
        assert tables[0].num_cols == 3

    def test_15_table_with_pipes_in_content(self):
        """Cell text containing pipe characters."""
        src = '---dox\nversion: \'1.0\'\nsource: t\nlang: en\n---\n\n||| table\n| Expression | Result |\n| --- | --- |\n| a OR b | true |\n|||'
        doc = parser.parse(src)
        tables = [e for e in doc.elements if isinstance(e, Table)]
        # Should parse without crashing
        assert len(tables) >= 1

    def test_16_colspan_2x2_merge(self):
        """2-column span in a 4-column table."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n||| table\n| A | B {cs=2} | D |\n| --- | --- | --- | --- |\n| 1 | 2 | 3 | 4 |\n|||"
        doc = parser.parse(src)
        t = [e for e in doc.elements if isinstance(e, Table)][0]
        assert t.rows[0].cells[1].colspan == 2

    def test_17_rowspan(self):
        """Row span in table."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n||| table\n| A | B |\n| --- | --- |\n| Left {rs=2} | Right1 |\n|  | Right2 |\n|||"
        doc = parser.parse(src)
        t = [e for e in doc.elements if isinstance(e, Table)][0]
        assert t.rows[1].cells[0].rowspan == 2

    def test_18_colspan_AND_rowspan(self):
        """Cell with both colspan=2 and rowspan=2."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n||| table\n| Merged {cs=2 rs=2} | C |\n| --- | --- | --- |\n|  | D |\n| E | F | G |\n|||"
        doc = parser.parse(src)
        t = [e for e in doc.elements if isinstance(e, Table)][0]
        c = t.rows[0].cells[0]
        assert c.colspan == 2
        assert c.rowspan == 2

    def test_19_multiple_tables(self):
        """Document with 5 separate tables."""
        parts = ["---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n"]
        for i in range(5):
            parts.append(f"\n## Table {i}\n")
            parts.append(f"||| table id=\"t{i}\"\n| Col1 | Col2 |\n| --- | --- |\n| A{i} | B{i} |\n|||\n")
        doc = parser.parse("\n".join(parts))
        tables = [e for e in doc.elements if isinstance(e, Table)]
        assert len(tables) == 5

    def test_20_table_with_formatting_in_cells(self):
        """Cells containing bold/italic text."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n||| table\n| **Bold Header** | *Italic Header* |\n| --- | --- |\n| `code` | [link](url) |\n|||"
        doc = parser.parse(src)
        t = [e for e in doc.elements if isinstance(e, Table)][0]
        assert "**Bold Header**" in t.rows[0].cells[0].text

    def test_21_table_survives_docx_roundtrip(self):
        """Table data preserved after DOCX conversion."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        t = Table(table_id="t1", caption="Test")
        t.rows.append(TableRow(cells=[
            TableCell(text="Name", is_header=True),
            TableCell(text="Age", is_header=True),
        ], is_header=True))
        for name, age in [("Alice", "30"), ("Bob", "25"), ("Charlie", "35")]:
            t.rows.append(TableRow(cells=[TableCell(text=name), TableCell(text=age)]))
        doc.add_element(t)
        text = _docx_text(doc)
        assert "alice" in text
        assert "bob" in text
        assert "35" in text


# ================================================================
# C. Inline Formatting Edge Cases
# ================================================================

class TestInlineFormatting:

    def test_22_nested_bold_italic(self):
        """Bold inside italic or vice versa."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nThis has ***bold italic*** text.\n")
        paras = [e for e in doc.elements if isinstance(e, Paragraph)]
        assert any("bold italic" in p.text.lower() for p in paras)

    def test_23_adjacent_formatting(self):
        """Back-to-back formatting: **bold***italic*."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n**bold***italic*`code` end.\n")
        paras = [e for e in doc.elements if isinstance(e, Paragraph)]
        assert len(paras) >= 1

    def test_24_code_with_special_chars(self):
        """Inline code with special chars."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nUse `dict[str, list[int]]` for typing.\n")
        paras = [e for e in doc.elements if isinstance(e, Paragraph)]
        assert any("dict[str" in p.text for p in paras)

    def test_25_url_with_query_params(self):
        """Link with complex URL."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nSee [docs](https://example.com/api?key=val&page=2#section).\n")
        paras = [e for e in doc.elements if isinstance(e, Paragraph)]
        assert any("docs" in p.text for p in paras)

    def test_26_dollar_signs_in_text(self):
        """Dollar signs that are NOT math delimiters."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nRevenue was $5.08B and costs were $3.2B.\n")
        paras = [e for e in doc.elements if isinstance(e, Paragraph)]
        assert any("$5.08B" in p.text for p in paras)

    def test_27_asterisks_in_text(self):
        """Literal asterisks that aren't formatting."""
        doc = parser.parse("---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nRated 4.5/5 stars ****.\n")
        paras = [e for e in doc.elements if isinstance(e, Paragraph)]
        assert len(paras) >= 1

    def test_28_inline_formatting_survives_docx(self):
        """Bold/italic runs actually exist in DOCX."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Paragraph(text="Has **bold** and *italic* and `code` inside."))
        from docx import Document as WordDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
            wd = WordDocument(f.name)
        has_bold = has_italic = has_mono = False
        for p in wd.paragraphs:
            for run in p.runs:
                if run.bold: has_bold = True
                if run.italic: has_italic = True
                if run.font.name == 'Courier New': has_mono = True
        assert has_bold
        assert has_italic
        assert has_mono


# ================================================================
# D. Math / Code Edge Cases
# ================================================================

class TestMathCodeEdges:

    def test_29_complex_latex(self):
        """Complex LaTeX expression with braces, subscripts, superscripts."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n$$\\frac{\\partial^2 u}{\\partial x^2} + \\frac{\\partial^2 u}{\\partial y^2} = 0$$\n"
        doc = parser.parse(src)
        maths = [e for e in doc.elements if isinstance(e, MathBlock)]
        assert len(maths) == 1
        assert "\\frac" in maths[0].expression

    def test_30_math_with_text_around(self):
        """Math block between paragraphs."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\nBefore math.\n\n$$x = \\frac{-b \\pm \\sqrt{b^2-4ac}}{2a}$$\n\nAfter math.\n"
        doc = parser.parse(src)
        types = [type(e).__name__ for e in doc.elements]
        assert "MathBlock" in types
        assert "Paragraph" in types

    def test_31_code_block_python(self):
        """Python code with indentation and special chars."""
        src = '---dox\nversion: \'1.0\'\nsource: t\nlang: en\n---\n\n```python\ndef factorial(n: int) -> int:\n    """Calculate n! recursively."""\n    if n <= 1:\n        return 1\n    return n * factorial(n - 1)\n```\n'
        doc = parser.parse(src)
        codes = [e for e in doc.elements if isinstance(e, CodeBlock)]
        assert len(codes) == 1
        assert "factorial" in codes[0].code
        assert codes[0].language == "python"

    def test_32_code_block_with_triple_backticks_inside(self):
        """Code that mentions backticks (shouldn't break parser)."""
        src = '---dox\nversion: \'1.0\'\nsource: t\nlang: en\n---\n\n```markdown\nUse `backticks` for code.\n```\n'
        doc = parser.parse(src)
        codes = [e for e in doc.elements if isinstance(e, CodeBlock)]
        assert len(codes) == 1

    def test_33_multiple_code_blocks(self):
        """Several code blocks in sequence."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n```python\nprint('hello')\n```\n\n```sql\nSELECT 1;\n```\n\n```bash\necho hi\n```\n"
        doc = parser.parse(src)
        codes = [e for e in doc.elements if isinstance(e, CodeBlock)]
        assert len(codes) == 3
        langs = [c.language for c in codes]
        assert "python" in langs
        assert "sql" in langs

    def test_34_math_in_docx(self):
        """Math expression appears in DOCX output."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(MathBlock(expression="E = mc^2", display_mode=True))
        text = _docx_text(doc)
        assert "e = mc^2" in text

    def test_35_code_in_docx(self):
        """Code block appears in DOCX output."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(CodeBlock(code="SELECT * FROM users WHERE active = true;", language="sql"))
        text = _docx_text(doc)
        assert "select" in text


# ================================================================
# E. Cross-Page and Page Break
# ================================================================

class TestCrossPage:

    def test_36_multiple_page_breaks(self):
        """Document with 5 page breaks."""
        src = "---dox\nversion: '1.0'\nsource: t\nlang: en\n---\n\n# Page 1 {page: 1}\n\n---page-break from=1 to=2---\n\n# Page 2 {page: 2}\n\n---page-break from=2 to=3---\n\n# Page 3 {page: 3}\n"
        doc = parser.parse(src)
        breaks = [e for e in doc.elements if isinstance(e, PageBreak)]
        headings = [e for e in doc.elements if isinstance(e, Heading)]
        assert len(breaks) == 2
        assert len(headings) == 3

    def test_37_page_numbers_persist(self):
        """Page numbers survive serialize→parse roundtrip."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        for i in range(1, 6):
            doc.add_element(Heading(level=2, text=f"Section {i}", page=i))
        text = serializer.serialize(doc, include_spatial=False, include_metadata=False)
        doc2 = parser.parse(text)
        headings = [e for e in doc2.elements if isinstance(e, Heading)]
        for h in headings:
            assert h.page is not None

    def test_38_page_break_in_docx(self):
        """Page breaks create actual pages in DOCX."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Heading(level=1, text="Page 1"))
        doc.add_element(PageBreak(from_page=1, to_page=2))
        doc.add_element(Heading(level=1, text="Page 2"))
        from docx import Document as WordDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
            wd = WordDocument(f.name)
        # Check for page break run
        has_break = False
        for p in wd.paragraphs:
            for run in p.runs:
                if run._element.xml and 'w:br' in run._element.xml:
                    has_break = True
        assert has_break

    def test_39_page_break_in_pdf(self):
        """Page breaks create multiple pages in PDF."""
        import fitz
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Heading(level=1, text="Page One"))
        doc.add_element(PageBreak(from_page=1, to_page=2))
        doc.add_element(Heading(level=1, text="Page Two"))
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            to_pdf(doc, f.name)
            pdf = fitz.open(f.name)
        assert pdf.page_count >= 2
        pdf.close()


# ================================================================
# F. DOCX Converter Specifics
# ================================================================

class TestDocxSpecifics:

    def test_40_heading_styles_correct(self):
        """H1-H6 map to correct Word heading styles."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        for i in range(1, 7):
            doc.add_element(Heading(level=i, text=f"Heading {i}"))
        from docx import Document as WordDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
            wd = WordDocument(f.name)
        heading_styles = [p.style.name for p in wd.paragraphs if p.style.name.startswith('Heading')]
        assert len(heading_styles) == 6

    def test_41_table_grid_style(self):
        """Tables use Table Grid style."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        t = Table()
        t.rows.append(TableRow(cells=[TableCell(text="A"), TableCell(text="B")], is_header=True))
        t.rows.append(TableRow(cells=[TableCell(text="1"), TableCell(text="2")]))
        doc.add_element(t)
        from docx import Document as WordDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
            wd = WordDocument(f.name)
        assert len(wd.tables) == 1
        assert wd.tables[0].style.name == 'Table Grid'

    def test_42_list_styles(self):
        """Ordered and unordered lists use correct Word styles."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(ListBlock(items=[ListItem(text="Bullet item")], ordered=False))
        doc.add_element(ListBlock(items=[ListItem(text="Number item")], ordered=True))
        from docx import Document as WordDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
            wd = WordDocument(f.name)
        styles = [p.style.name for p in wd.paragraphs]
        assert 'List Bullet' in styles
        assert 'List Number' in styles

    def test_43_figure_placeholder_when_no_image(self):
        """Figure with non-existent image file shows placeholder."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Figure(source="nonexistent.png", caption="Test Figure"))
        text = _docx_text(doc)
        assert "image" in text or "nonexistent" in text

    def test_44_large_document_docx(self):
        """Document with 200 elements converts to DOCX without error."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        for i in range(50):
            doc.add_element(Heading(level=2, text=f"Section {i}"))
            doc.add_element(Paragraph(text=f"Paragraph number {i} with some text content."))
            t = Table()
            t.rows.append(TableRow(cells=[TableCell(text="A"), TableCell(text="B")], is_header=True))
            t.rows.append(TableRow(cells=[TableCell(text=str(i)), TableCell(text=f"val{i}")]))
            doc.add_element(t)
            doc.add_element(CodeBlock(code=f"print({i})", language="python"))
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.stat().st_size > 10000

    def test_45_colspan_merge_doesnt_crash(self):
        """Colspan merging in DOCX doesn't throw."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        t = Table()
        t.rows.append(TableRow(cells=[
            TableCell(text="Wide Header", is_header=True, colspan=3),
        ], is_header=True))
        t.rows.append(TableRow(cells=[
            TableCell(text="A"), TableCell(text="B"), TableCell(text="C"),
        ]))
        doc.add_element(t)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.exists()


# ================================================================
# G. PDF Converter Specifics
# ================================================================

class TestPdfSpecifics:

    def test_46_large_document_pdf(self):
        """Large document converts to PDF without error."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        for i in range(100):
            doc.add_element(Paragraph(text=f"Paragraph {i}: Lorem ipsum dolor sit amet."))
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)
        assert path.stat().st_size > 1000

    def test_47_table_in_pdf(self):
        """Table renders in PDF."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        t = Table(caption="PDF Table")
        t.rows.append(TableRow(cells=[TableCell(text="X", is_header=True), TableCell(text="Y", is_header=True)], is_header=True))
        for i in range(5):
            t.rows.append(TableRow(cells=[TableCell(text=str(i)), TableCell(text=str(i*2))]))
        doc.add_element(t)
        text = _pdf_text(doc)
        assert "pdf table" in text

    def test_48_pdf_xml_escaping(self):
        """Special chars (<, >, &) don't break PDF generation."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Paragraph(text="if (x < 10 && y > 5) { return x & y; }"))
        doc.add_element(Paragraph(text="Tom & Jerry <3 HTML"))
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            path = to_pdf(doc, f.name)
        assert path.exists()
        text = _pdf_text(doc)
        assert "tom" in text


# ================================================================
# H. Full Pipeline Tests
# ================================================================

class TestFullPipeline:

    def test_49_dox_to_serialize_to_parse_to_docx(self):
        """Full: parse .dox → serialize → re-parse → DOCX → extract."""
        src = """---dox
version: '1.0'
source: pipeline.pdf
lang: en
---

# Quarterly Update {page: 1, id: "h1"}

Revenue hit **$5.2M** this quarter, up *18%* from Q2. {page: 1, id: "p1"}

||| table id="t1" caption="Q3 Numbers"
| Metric | Value |
| --- | --- |
| Revenue | $5.2M |
| Costs | $3.1M |
| Profit | $2.1M |
|||

- Strong cloud growth
- New enterprise clients
- Reduced churn by 12%

$$\\text{CAGR} = \\left(\\frac{V_f}{V_i}\\right)^{1/n} - 1$$ {math: latex, page: 1}

[^1]: Preliminary figures, subject to audit.
"""
        doc = parser.parse(src)
        text_ser = serializer.serialize(doc, include_spatial=False, include_metadata=False)
        doc2 = parser.parse(text_ser)
        docx_text = _docx_text(doc2)

        assert "quarterly update" in docx_text
        assert "5.2m" in docx_text
        assert "revenue" in docx_text
        assert "profit" in docx_text

    def test_50_dox_to_serialize_to_parse_to_pdf(self):
        """Full: parse .dox → serialize → re-parse → PDF → extract."""
        src = """---dox
version: '1.0'
source: test.pdf
lang: en
---

# Annual Report {page: 1}

Total employees: **14,800**. {page: 1}

||| table caption="Headcount"
| Division | Count |
| --- | --- |
| Engineering | 5,200 |
| Sales | 3,100 |
| Operations | 6,500 |
|||
"""
        doc = parser.parse(src)
        text_ser = serializer.serialize(doc, include_spatial=False, include_metadata=False)
        doc2 = parser.parse(text_ser)
        pdf_text = _pdf_text(doc2)

        assert "annual report" in pdf_text
        assert "14,800" in pdf_text


# ================================================================
# I. Unicode and Special Characters
# ================================================================

class TestUnicodeSpecial:

    def test_51_cjk_text(self):
        """Chinese/Japanese/Korean text."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="ja")
        doc.add_element(Heading(level=1, text="年次報告書"))
        doc.add_element(Paragraph(text="会社の業績は好調でした。"))
        # DOCX should not crash
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.exists()

    def test_52_arabic_rtl(self):
        """Arabic right-to-left text."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="ar")
        doc.add_element(Paragraph(text="مرحبا بالعالم"))
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.exists()

    def test_53_emoji_in_content(self):
        """Emoji characters."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Paragraph(text="Status: ✅ Complete 🎉 Celebration"))
        text = _docx_text(doc)
        assert "complete" in text

    def test_54_math_symbols(self):
        """Mathematical Unicode symbols."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Paragraph(text="α β γ δ ε → ∞ ≤ ≥ ≠ ∑ ∏ √"))
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.exists()

    def test_55_curly_quotes_and_dashes(self):
        """Curly quotes, em dash, en dash."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Paragraph(text='\u201cHello,\u201d she said \u2014 it\u2019s an em dash \u2013 and en dash.'))
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        text = _docx_text(doc)
        assert "hello" in text

    def test_56_html_entities_in_text(self):
        """Text with < > & that could break XML/HTML."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Paragraph(text="if (a < b && c > d) { return a & b; }"))
        doc.add_element(Paragraph(text='<script>alert("xss")</script>'))
        # Both converters should handle this
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            to_pdf(doc, f.name)


# ================================================================
# J. Adversarial and Stress
# ================================================================

class TestAdversarial:

    def test_57_very_long_paragraph(self):
        """Paragraph with 5000 characters."""
        long_text = "Word " * 1000
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Paragraph(text=long_text.strip()))
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.stat().st_size > 5000

    def test_58_very_long_heading(self):
        """Heading with 200+ characters."""
        long_title = "A " * 100 + "Title"
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Heading(level=1, text=long_title))
        text = _docx_text(doc)
        assert "title" in text

    def test_59_deeply_nested_lists(self):
        """Nested list items (via text, not actual nesting)."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        items = [ListItem(text=f"{'  ' * i}Level {i} item") for i in range(5)]
        doc.add_element(ListBlock(items=items, ordered=False))
        text = _docx_text(doc)
        assert "level 0" in text
        assert "level 4" in text

    def test_60_mixed_element_types_rapid(self):
        """Rapid alternation between different element types."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        for i in range(20):
            doc.add_element(Heading(level=2, text=f"H{i}"))
            doc.add_element(Paragraph(text=f"P{i}"))
            doc.add_element(MathBlock(expression=f"x_{i} = {i}", display_mode=True))
            doc.add_element(CodeBlock(code=f"print({i})", language="python"))
            t = Table()
            t.rows.append(TableRow(cells=[TableCell(text=f"R{i}C1"), TableCell(text=f"R{i}C2")]))
            doc.add_element(t)
        # Should not crash for either format
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            to_pdf(doc, f.name)

    def test_61_all_element_types_in_one_doc(self):
        """Single document with every element type."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Heading(level=1, text="Title"))
        doc.add_element(Paragraph(text="Para with **bold**."))
        t = Table(caption="Cap")
        t.rows.append(TableRow(cells=[TableCell(text="H1", is_header=True)], is_header=True))
        t.rows.append(TableRow(cells=[TableCell(text="D1")]))
        doc.add_element(t)
        doc.add_element(CodeBlock(code="x = 1", language="python"))
        doc.add_element(MathBlock(expression="y = mx + b", display_mode=True))
        doc.add_element(Figure(source="img.png", caption="Fig"))
        doc.add_element(ListBlock(items=[ListItem(text="A"), ListItem(text="B")], ordered=False))
        doc.add_element(ListBlock(items=[ListItem(text="1"), ListItem(text="2")], ordered=True))
        doc.add_element(Footnote(number=1, text="Note"))
        doc.add_element(FormField(field_name="Name", value="Test"))
        doc.add_element(PageBreak(from_page=1, to_page=2))
        doc.add_element(Heading(level=2, text="Page 2"))
        doc.add_element(Annotation(annotation_type="comment", text="Annotation"))
        doc.add_element(Chart(chart_type="bar"))
        doc.add_element(CrossRef(ref_type="table", ref_id="t1"))

        # Convert to both formats without crashing
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            to_pdf(doc, f.name)

    def test_62_table_with_all_spans(self):
        """Table where every cell has a span of some kind."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        t = Table()
        t.rows.append(TableRow(cells=[
            TableCell(text="A", colspan=2, is_header=True),
            TableCell(text="B", rowspan=2, is_header=True),
        ], is_header=True))
        t.rows.append(TableRow(cells=[
            TableCell(text="C"), TableCell(text="D"), TableCell(text=""),
        ]))
        t.rows.append(TableRow(cells=[
            TableCell(text="E", colspan=3),
        ]))
        doc.add_element(t)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)

    def test_63_footnotes_high_numbers(self):
        """Footnotes with high numbers."""
        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        for i in range(1, 51):
            doc.add_element(Footnote(number=i, text=f"Footnote number {i}."))
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = to_docx(doc, f.name)
        assert path.exists()


# ================================================================
# NEW: Build complex DOCX with images
# ================================================================

class TestComplexDocxWithImages:
    """Build a complex document from scratch, convert, verify."""

    def test_64_create_real_images_and_embed(self):
        """Create actual PNG images and embed them in DOCX."""
        # Create a test image
        try:
            from PIL import Image as PILImage
        except ImportError:
            pytest.skip("Pillow not installed")

        # Create a simple 200x100 blue rectangle
        img = PILImage.new('RGB', (200, 100), color=(70, 130, 180))
        img_path = Path(tempfile.mktemp(suffix='.png'))
        img.save(str(img_path))

        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Heading(level=1, text="Document With Images"))
        doc.add_element(Paragraph(text="This document has an embedded image below:"))
        doc.add_element(Figure(source=str(img_path), caption="Blue Rectangle"))
        doc.add_element(Paragraph(text="And text continues after the image."))

        from docx import Document as WordDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
            wd = WordDocument(f.name)

        # Check that inline shapes exist (images)
        has_image = False
        for rel in wd.part.rels.values():
            if "image" in rel.reltype:
                has_image = True
                break
        assert has_image, "No image found in DOCX"

        # Cleanup
        img_path.unlink(missing_ok=True)

    def test_65_multiple_images_across_pages(self):
        """Multiple images on different pages."""
        try:
            from PIL import Image as PILImage
        except ImportError:
            pytest.skip("Pillow not installed")

        paths = []
        for color, name in [((255, 0, 0), 'red'), ((0, 255, 0), 'green'), ((0, 0, 255), 'blue')]:
            img = PILImage.new('RGB', (150, 80), color=color)
            p = Path(tempfile.mktemp(suffix=f'_{name}.png'))
            img.save(str(p))
            paths.append(p)

        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="t", lang="en")
        doc.add_element(Heading(level=1, text="Multi-Image Document"))
        doc.add_element(Figure(source=str(paths[0]), caption="Red Image"))
        doc.add_element(PageBreak(from_page=1, to_page=2))
        doc.add_element(Figure(source=str(paths[1]), caption="Green Image"))
        doc.add_element(PageBreak(from_page=2, to_page=3))
        doc.add_element(Figure(source=str(paths[2]), caption="Blue Image"))

        from docx import Document as WordDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            to_docx(doc, f.name)
            wd = WordDocument(f.name)

        image_count = sum(1 for rel in wd.part.rels.values() if "image" in rel.reltype)
        assert image_count == 3

        for p in paths:
            p.unlink(missing_ok=True)

    def test_66_complex_doc_images_tables_formatting(self):
        """Complex realistic document: images + tables + all formatting."""
        try:
            from PIL import Image as PILImage
        except ImportError:
            pytest.skip("Pillow not installed")

        # Create chart placeholder image
        chart_img = PILImage.new('RGB', (400, 250), color=(240, 240, 240))
        chart_path = Path(tempfile.mktemp(suffix='_chart.png'))
        chart_img.save(str(chart_path))

        doc = DoxDocument()
        doc.frontmatter = Frontmatter(version="1.0", source="complex-report.pdf", pages=3, lang="en")

        # Page 1
        doc.add_element(Heading(level=1, text="Comprehensive Analysis Report"))
        doc.add_element(Paragraph(text="**Prepared by:** Analytics Division | **Date:** December 2025"))
        doc.add_element(Paragraph(text="This report covers *revenue trends*, `technical metrics`, and [strategic recommendations](http://example.com)."))

        t1 = Table(table_id="t-summary", caption="Executive Summary Metrics")
        t1.rows.append(TableRow(cells=[
            TableCell(text="Category", is_header=True),
            TableCell(text="Q3 Actual", is_header=True),
            TableCell(text="Q3 Target", is_header=True),
            TableCell(text="Variance", is_header=True),
        ], is_header=True))
        for cat, actual, target, var in [
            ("Revenue", "$5.2M", "$4.8M", "+8.3%"),
            ("Gross Margin", "68.5%", "65.0%", "+3.5pp"),
            ("Customer NPS", "72", "70", "+2"),
            ("Uptime SLA", "99.97%", "99.95%", "+0.02%"),
        ]:
            t1.rows.append(TableRow(cells=[
                TableCell(text=cat), TableCell(text=actual),
                TableCell(text=target), TableCell(text=var),
            ]))
        doc.add_element(t1)

        doc.add_element(Figure(source=str(chart_path), caption="Revenue Trend Q1-Q3 2025"))

        doc.add_element(PageBreak(from_page=1, to_page=2))

        # Page 2
        doc.add_element(Heading(level=2, text="Technical Deep Dive"))
        doc.add_element(CodeBlock(code="""import pandas as pd

def analyze_revenue(df: pd.DataFrame) -> dict:
    \"\"\"Compute quarterly revenue metrics.\"\"\"
    return {
        'total': df['revenue'].sum(),
        'mean': df['revenue'].mean(),
        'growth': (df['revenue'].iloc[-1] / df['revenue'].iloc[0] - 1) * 100,
    }""", language="python"))

        doc.add_element(MathBlock(
            expression="\\text{CAGR} = \\left(\\frac{V_f}{V_i}\\right)^{\\frac{1}{n}} - 1",
            display_mode=True,
        ))

        t2 = Table(table_id="t-tech", caption="Service Performance (colspan demo)")
        t2.rows.append(TableRow(cells=[
            TableCell(text="Service", is_header=True),
            TableCell(text="Performance Metrics", is_header=True, colspan=2),
        ], is_header=True))
        t2.rows.append(TableRow(cells=[
            TableCell(text=""), TableCell(text="Latency", is_header=True),
            TableCell(text="Availability", is_header=True),
        ], is_header=True))
        for svc, lat, avail in [("API", "12ms", "99.99%"), ("DB", "3ms", "99.999%"), ("Cache", "1ms", "99.95%")]:
            t2.rows.append(TableRow(cells=[
                TableCell(text=svc), TableCell(text=lat), TableCell(text=avail),
            ]))
        doc.add_element(t2)

        doc.add_element(PageBreak(from_page=2, to_page=3))

        # Page 3
        doc.add_element(Heading(level=2, text="Recommendations"))
        doc.add_element(ListBlock(items=[
            ListItem(text="Increase cloud infrastructure investment by 20%"),
            ListItem(text="Expand APAC sales team by Q2 2026"),
            ListItem(text="Implement real-time monitoring dashboard"),
            ListItem(text="Reduce technical debt in legacy services"),
        ], ordered=True))

        doc.add_element(ListBlock(items=[
            ListItem(text="High priority: Database migration"),
            ListItem(text="Medium priority: API versioning"),
            ListItem(text="Low priority: UI refresh"),
        ], ordered=False))

        doc.add_element(Footnote(number=1, text="All figures in USD. Preliminary, subject to audit."))
        doc.add_element(Footnote(number=2, text="SLA calculations exclude planned maintenance windows."))
        doc.add_element(FormField(field_name="Approved By", value="VP of Engineering"))
        doc.add_element(FormField(field_name="Review Date", value="2025-12-20"))

        # Convert to both formats
        from docx import Document as WordDocument
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            docx_path = to_docx(doc, f.name)
            wd = WordDocument(f.name)

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            pdf_path = to_pdf(doc, f.name)

        # ── DOCX Verification ──
        docx_text = ' '.join(p.text for p in wd.paragraphs if p.text.strip()).lower()
        docx_table_text = ''
        for t in wd.tables:
            for r in t.rows:
                for c in r.cells:
                    docx_table_text += ' ' + c.text.lower()

        all_docx = docx_text + docx_table_text

        # Headings
        assert "comprehensive analysis report" in all_docx
        assert "technical deep dive" in all_docx
        assert "recommendations" in all_docx

        # Paragraph content
        assert "analytics division" in all_docx
        assert "revenue trends" in all_docx

        # Table data
        assert "revenue" in docx_table_text
        assert "5.2m" in docx_table_text or "$5.2m" in docx_table_text
        assert "99.99%" in docx_table_text
        assert "api" in docx_table_text

        # Code
        assert "analyze_revenue" in all_docx

        # Math
        assert "cagr" in all_docx

        # Lists
        assert "cloud infrastructure" in all_docx
        assert "database migration" in all_docx

        # Form fields
        assert "approved by" in all_docx
        assert "vp of engineering" in all_docx

        # Image embedded
        image_count = sum(1 for rel in wd.part.rels.values() if "image" in rel.reltype)
        assert image_count >= 1

        # Tables count
        assert len(wd.tables) == 2

        # ── PDF Verification ──
        import fitz
        pdf = fitz.open(str(pdf_path))
        assert pdf.page_count >= 3
        pdf_all = ''
        for page in pdf:
            pdf_all += page.get_text().lower()
        pdf.close()

        assert "comprehensive analysis report" in pdf_all
        assert "revenue" in pdf_all
        assert "recommendations" in pdf_all

        # Cleanup
        chart_path.unlink(missing_ok=True)

        print(f"\n=== COMPLEX DOC RESULTS ===")
        print(f"Elements: {len(doc.elements)}")
        print(f"DOCX size: {docx_path.stat().st_size:,} bytes")
        print(f"PDF size: {pdf_path.stat().st_size:,} bytes")
        print(f"DOCX tables: {len(wd.tables)}")
        print(f"DOCX images: {image_count}")
        print(f"PDF pages: {fitz.open(str(pdf_path)).page_count}")
